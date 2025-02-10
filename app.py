import os
import io
import subprocess
import shutil

from flask import Flask, render_template, request, redirect, url_for, send_file, flash
from werkzeug.utils import secure_filename

# ================
# = Librairies PDF
# ================
import PyPDF2
from PyPDF2 import PdfReader, PdfWriter, PdfMerger

from pdf2docx import Converter
from pdf2image import convert_from_path

import fitz  # PyMuPDF
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

from docx import Document
from reportlab.lib.utils import simpleSplit


# ================
# = Config & Setup
# ================
os.environ["GIO_USE_VFS"] = "local"  # Désactive certains warnings GIO
app = Flask(__name__)
app.secret_key = "secret_key_flask"  # Pour utiliser flash()

UPLOAD_FOLDER = os.path.join(os.getcwd(), 'uploads')
OUTPUT_FOLDER = os.path.join(os.getcwd(), 'outputs')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER


# ================
# = Fonctions PDF
# ================

def is_pdf_file(filename: str) -> bool:
    """
    Vérifie si le fichier possède l'extension .pdf (insensible à la casse).
    """
    return '.' in filename and filename.rsplit('.', 1)[1].lower() == 'pdf'


def merge_pdfs(input_files, output_file):
    pdf_merger = PdfMerger()
    for pdf in input_files:
        pdf_merger.append(pdf)
    pdf_merger.write(output_file)
    pdf_merger.close()

def split_pdf(input_file, output_dir, pages_per_split=None, page_range=None):
    pdf_reader = PdfReader(input_file)

    if page_range:
        start, end = map(int, page_range.split('-'))
        pdf_writer = PdfWriter()
        for page_num in range(start - 1, end):
            pdf_writer.add_page(pdf_reader.pages[page_num])
        output_filename = os.path.join(output_dir, f'pages_{start}_to_{end}.pdf')
        with open(output_filename, 'wb') as out:
            pdf_writer.write(out)
        return [output_filename]

    elif pages_per_split:
        output_files = []
        pages_per_split = int(pages_per_split)
        total_pages = len(pdf_reader.pages)
        for i in range(0, total_pages, pages_per_split):
            pdf_writer = PdfWriter()
            for j in range(i, min(i + pages_per_split, total_pages)):
                pdf_writer.add_page(pdf_reader.pages[j])
            output_filename = os.path.join(
                output_dir,
                f'pages_{i + 1}_to_{min(i + pages_per_split, total_pages)}.pdf'
            )
            with open(output_filename, 'wb') as out:
                pdf_writer.write(out)
            output_files.append(output_filename)
        return output_files

    else:
        # Découpe page par page
        output_files = []
        for page_num in range(len(pdf_reader.pages)):
            pdf_writer = PdfWriter()
            pdf_writer.add_page(pdf_reader.pages[page_num])
            output_filename = os.path.join(output_dir, f'page_{page_num + 1}.pdf')
            with open(output_filename, 'wb') as out:
                pdf_writer.write(out)
            output_files.append(output_filename)
        return output_files


# Ajouter Ghostscript au PATH si nécessaire
gs_path = r"C:\Program Files\gs\gs10.04.0\bin"
os.environ["PATH"] += os.pathsep + gs_path

def compress_pdf(input_file, output_file, quality="medium"):
    """
    Compresse un fichier PDF en réduisant la taille des images et en optimisant avec Ghostscript.

    quality peut être :
    - "low" : Compression maximale (images très compressées, qualité faible)
    - "medium" : Bon compromis entre compression et qualité
    - "high" : Compression faible (qualité élevée)
    """
    print(f"🔧 Début de la compression : {input_file} → {output_file}")

    # Mapping des niveaux de compression Ghostscript
    gs_quality_map = {
        "low": "/screen",
        "medium": "/ebook",
        "high": "/printer"
    }
    gs_quality = gs_quality_map.get(quality, "/ebook")  # Par défaut : ebook

    # Étape 1 : Réduction de la taille des images
    temp_pdf = input_file.replace(".pdf", "_temp.pdf")
    compress_images_in_pdf(input_file, temp_pdf, dpi=100, quality=70)

    # Étape 2 : Optimisation avec Ghostscript
    compress_with_ghostscript(temp_pdf, output_file, gs_quality)

    # Supprimer le fichier temporaire
    if os.path.exists(temp_pdf):
        os.remove(temp_pdf)

    print(f"✅ Compression terminée : {output_file}")


def compress_images_in_pdf(input_file, output_file, dpi=100, quality=70):
    """
    Réduit la résolution et la qualité des images dans un PDF.
    """
    print(f"📉 Réduction des images : DPI={dpi}, qualité={quality}")

    doc = fitz.open(input_file)

    for page_num in range(len(doc)):
        images = doc[page_num].get_images(full=True)
        for img_index, img in enumerate(images):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            img_ext = base_image["ext"]
            img_pil = Image.open(io.BytesIO(image_bytes))

            # Convertir en JPEG (réduction qualité) si ce n'est pas déjà du JPEG
            if img_ext.lower() != "jpeg":
                img_pil = img_pil.convert("RGB")

            # ✅ Correction ici : Utilisation de `Image.Resampling.LANCZOS` au lieu de `Image.ANTIALIAS`
            img_pil = img_pil.resize((img_pil.width // 2, img_pil.height // 2), Image.Resampling.LANCZOS)
            
            img_bytes_io = io.BytesIO()
            img_pil.save(img_bytes_io, format="JPEG", quality=quality)
            
            # Remplacer l'image originale par l'image compressée
            doc[page_num].insert_image(doc[page_num].rect, stream=img_bytes_io.getvalue(), xref=xref)

    # Sauvegarde du PDF modifié
    doc.save(output_file)
    doc.close()
    print(f"📜 PDF avec images compressées sauvegardé : {output_file}")


def compress_with_ghostscript(input_file, output_file, gs_quality="/ebook"):
    """
    Utilise Ghostscript pour compresser un PDF entier.
    """
    gs_executable = "gswin64c" if os.name == "nt" else "gs"

    if shutil.which(gs_executable) is None:
        raise RuntimeError("❌ Ghostscript n'est pas installé ou introuvable !")

    gs_command = [
        gs_executable,
        "-sDEVICE=pdfwrite",
        "-dCompatibilityLevel=1.4",
        "-dNOPAUSE",
        "-dQUIET",
        "-dBATCH",
        f"-dPDFSETTINGS={gs_quality}",
        f"-sOutputFile={output_file}",
        input_file
    ]

    try:
        print(f"🔄 Exécution de Ghostscript : {' '.join(gs_command)}")
        subprocess.run(gs_command, check=True)
        print(f"✅ PDF optimisé avec Ghostscript : {output_file}")
    except subprocess.CalledProcessError as e:
        print(f"❌ Erreur Ghostscript : {e}")
        raise RuntimeError(f"Erreur Ghostscript : {e}")



def pdf_to_images(input_file, output_dir):
    images = convert_from_path(input_file)
    image_paths = []
    for i, img in enumerate(images):
        image_path = os.path.join(output_dir, f'page_{i+1}.png')
        img.save(image_path, 'PNG')
        image_paths.append(image_path)
    return image_paths

def pdf_to_word(input_file, output_file):
    from docx import Document
    from docx.shared import Inches
    try:
        print(f"🔄 Conversion PDF→Word : {input_file} → {output_file}")

        doc = Document()
        pdf_document = fitz.open(input_file)

        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            text = page.get_text("text")

            if text.strip():
                doc.add_paragraph(f"📄 Page {page_num + 1} :")
                doc.add_paragraph(text)
            else:
                doc.add_paragraph(f"📄 Page {page_num + 1} : (Aucun texte détecté)")

        doc.save(output_file)
        print(f"✅ Conversion PDF→Word terminée : {output_file}")

    except Exception as e:
        print(f"❌ Erreur conversion PDF→Word : {e}")


def rotate_pdf(input_file, output_file, angle):
    pdf_reader = PdfReader(input_file)
    pdf_writer = PdfWriter()
    for page in pdf_reader.pages:
        page.rotate(angle)
        pdf_writer.add_page(page)
    with open(output_file, 'wb') as out:
        pdf_writer.write(out)

def add_page_numbers(input_file, output_file):
    pdf_reader = PdfReader(input_file)
    pdf_writer = PdfWriter()

    for i, page in enumerate(pdf_reader.pages):
        packet = io.BytesIO()
        can = canvas.Canvas(packet, pagesize=A4)
        can.drawString(500, 10, str(i+1))
        can.save()

        packet.seek(0)
        watermark = PdfReader(packet).pages[0]
        page.merge_page(watermark)
        pdf_writer.add_page(page)
    with open(output_file, 'wb') as out:
        pdf_writer.write(out)

def repair_pdf(input_file, output_file):
    reader = PyPDF2.PdfReader(input_file)
    writer = PyPDF2.PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    with open(output_file, 'wb') as out:
        writer.write(out)
    print(f"✅ PDF réparé : {output_file}")

def unlock_pdf(input_file, output_file, password):
    pdf_reader = PyPDF2.PdfReader(input_file)
    if pdf_reader.is_encrypted:
        pdf_reader.decrypt(password)
    pdf_writer = PyPDF2.PdfWriter()
    for page in pdf_reader.pages:
        pdf_writer.add_page(page)
    with open(output_file, 'wb') as out:
        pdf_writer.write(out)

def protect_pdf(input_file, output_file, password):
    pdf_reader = PyPDF2.PdfReader(input_file)
    pdf_writer = PyPDF2.PdfWriter()
    for page in pdf_reader.pages:
        pdf_writer.add_page(page)
    pdf_writer.encrypt(password)
    with open(output_file, 'wb') as out:
        pdf_writer.write(out)

def images_to_pdf(images, output_file):
    """
    Convertit une liste d'images en un fichier PDF.
    """
    img_list = []
    for img in images:
        image = Image.open(img)
        image = image.convert('RGB')  # Convertit en mode RGB pour éviter les problèmes
        img_list.append(image)

    if img_list:
        img_list[0].save(output_file, save_all=True, append_images=img_list[1:])

def word_to_pdf(input_file, output_file):
    """
    Convertit un fichier .docx en PDF avec gestion automatique du Word Wrap et des marges.
    """
    try:
        # Charger le document Word
        doc = Document(input_file)

        # Définir le PDF
        c = canvas.Canvas(output_file, pagesize=A4)
        width, height = A4

        # Définition des marges
        margin_left = 50
        margin_right = 50
        margin_top = height - 50
        margin_bottom = 50

        max_width = width - margin_left - margin_right
        y_position = margin_top
        line_height = 20  # Espacement des lignes

        for para in doc.paragraphs:
            if para.text.strip():  # Ignorer les lignes vides
                wrapped_text = simpleSplit(para.text, "Helvetica", 12, max_width)

                for line in wrapped_text:
                    if y_position < margin_bottom:
                        c.showPage()  # Nouvelle page
                        y_position = margin_top

                    c.drawString(margin_left, y_position, line)
                    y_position -= line_height

        c.save()  # Sauvegarde le PDF
        print(f"✅ Conversion réussie : {output_file}")

    except Exception as e:
        raise RuntimeError(f"❌ Erreur lors de la conversion : {e}")


# =========================
# = Routes Flask
# =========================

@app.route('/home')
def index():
    """ Page d'accueil avec formulaires """
    return render_template('pdf_ascend_tool.html')

@app.route('/')
def loading():
    """
    Page de chargement animée avec redirection automatique.
    """
    return render_template('loading.html', redirect_url=url_for('index'))

@app.route('/merge', methods=['POST'])
def route_merge():
    files = request.files.getlist('merge_files')
    if not files or len(files) == 0:
        flash("Aucun fichier sélectionné pour la fusion.")
        return redirect(url_for('index'))

    saved_paths = []
    for f in files:
        if f.filename.strip() == '':
            continue
        if not is_pdf_file(f.filename):
            flash(f"Fichier invalide '{f.filename}' : veuillez sélectionner un PDF.")
            return redirect(url_for('index'))

        filename = secure_filename(f.filename)
        f_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        f.save(f_path)
        saved_paths.append(f_path)

    if not saved_paths:
        flash("Veuillez sélectionner au moins un fichier PDF.")
        return redirect(url_for('index'))

    output_file = os.path.join(app.config['OUTPUT_FOLDER'], 'merged.pdf')
    merge_pdfs(saved_paths, output_file)

    flash("Fusion terminée.")
    return send_file(output_file, as_attachment=True)


@app.route('/split', methods=['POST'])
def route_split():
    file = request.files.get('split_file')
    pages_per_split = request.form.get('pages_per_split')
    page_range = request.form.get('page_range')

    if not file or file.filename.strip() == '':
        flash("Aucun fichier sélectionné pour le split.")
        return redirect(url_for('index'))

    if not is_pdf_file(file.filename):
        flash(f"Fichier invalide '{file.filename}' : veuillez sélectionner un PDF.")
        return redirect(url_for('index'))

    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)

    output_files = split_pdf(file_path, app.config['OUTPUT_FOLDER'], pages_per_split, page_range)
    if len(output_files) == 1:
        flash("Split terminé (un seul fichier).")
        return send_file(output_files[0], as_attachment=True)
    else:
        import zipfile
        zip_filename = os.path.join(app.config['OUTPUT_FOLDER'], 'split_files.zip')
        with zipfile.ZipFile(zip_filename, 'w') as zipf:
            for f in output_files:
                zipf.write(f, arcname=os.path.basename(f))
        flash("Split terminé (plusieurs fichiers).")
        return send_file(zip_filename, as_attachment=True)


@app.route('/compress', methods=['POST'])
def route_compress():
    file = request.files.get('compress_file')
    compression_quality = request.form.get('compression_quality')

    if not file or file.filename.strip() == '':
        flash("⚠️ Aucun fichier sélectionné pour la compression.")
        return redirect(url_for('index'))

    if not is_pdf_file(file.filename):
        flash(f"Fichier invalide '{file.filename}' : veuillez sélectionner un PDF.")
        return redirect(url_for('index'))

    if not compression_quality:
        flash("⚠️ Qualité de compression manquante.")
        return redirect(url_for('index'))

    valid_qualities = ["low", "medium", "high"]
    if compression_quality not in valid_qualities:
        flash(f"⚠️ Qualité invalide ! Choisissez parmi : {', '.join(valid_qualities)}")
        return redirect(url_for('index'))

    filename = secure_filename(file.filename)
    in_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(in_path)

    out_path = os.path.join(app.config['OUTPUT_FOLDER'], 'compressed_' + filename)
    try:
        compress_pdf(in_path, out_path, compression_quality)
        flash(f"✅ Compression réussie ({compression_quality})")
        return send_file(out_path, as_attachment=True)
    except Exception as e:
        flash(f"❌ Erreur compression : {e}")
        return redirect(url_for('index'))



@app.route('/pdf_to_images', methods=['POST'])
def route_pdf_to_images():
    file = request.files.get('images_file')
    if not file or file.filename.strip() == '':
        flash("Aucun fichier sélectionné pour PDF → Images.")
        return redirect(url_for('index'))

    if not is_pdf_file(file.filename):
        flash(f"Fichier invalide '{file.filename}' : veuillez sélectionner un PDF.")
        return redirect(url_for('index'))

    filename = secure_filename(file.filename)
    in_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(in_path)

    images = pdf_to_images(in_path, app.config['OUTPUT_FOLDER'])
    if images:
        import zipfile
        zip_filename = os.path.join(app.config['OUTPUT_FOLDER'], 'images.zip')
        with zipfile.ZipFile(zip_filename, 'w') as zipf:
            for img_path in images:
                zipf.write(img_path, arcname=os.path.basename(img_path))
        flash("Conversion PDF → Images terminée.")
        return send_file(zip_filename, as_attachment=True)
    else:
        flash("Impossible de convertir le PDF en images.")
        return redirect(url_for('index'))


@app.route('/pdf_to_word', methods=['POST'])
def route_pdf_to_word():
    file = request.files.get('word_file')
    if not file or file.filename.strip() == '':
        flash("Aucun fichier sélectionné pour PDF → Word.")
        return redirect(url_for('index'))

    if not is_pdf_file(file.filename):
        flash(f"Fichier invalide '{file.filename}' : veuillez sélectionner un PDF.")
        return redirect(url_for('index'))

    filename = secure_filename(file.filename)
    in_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(in_path)

    out_path = os.path.join(app.config['OUTPUT_FOLDER'], filename.replace(".pdf", ".docx"))
    try:
        pdf_to_word(in_path, out_path)
        flash("Conversion PDF → Word terminée.")
        return send_file(out_path, as_attachment=True)
    except Exception as e:
        flash(f"Erreur conversion PDF → Word : {e}")
        return redirect(url_for('index'))


@app.route('/rotate', methods=['POST'])
def route_rotate():
    file = request.files.get('rotate_file')
    angle = request.form.get('angle')

    if not file or file.filename.strip() == '':
        flash("Aucun fichier sélectionné pour la rotation.")
        return redirect(url_for('index'))
    if not angle:
        flash("Angle de rotation manquant.")
        return redirect(url_for('index'))

    if not is_pdf_file(file.filename):
        flash(f"Fichier invalide '{file.filename}' : veuillez sélectionner un PDF.")
        return redirect(url_for('index'))

    filename = secure_filename(file.filename)
    in_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(in_path)

    out_path = os.path.join(app.config['OUTPUT_FOLDER'], 'rotated.pdf')
    rotate_pdf(in_path, out_path, int(angle))

    flash("Rotation terminée.")
    return send_file(out_path, as_attachment=True)


@app.route('/add_page_numbers', methods=['POST'])
def route_add_page_numbers():
    file = request.files.get('numbers_file')
    if not file or file.filename.strip() == '':
        flash("Aucun fichier PDF pour l'ajout de numéros de page.")
        return redirect(url_for('index'))

    if not is_pdf_file(file.filename):
        flash(f"Fichier invalide '{file.filename}' : veuillez sélectionner un PDF.")
        return redirect(url_for('index'))

    filename = secure_filename(file.filename)
    in_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(in_path)

    out_path = os.path.join(app.config['OUTPUT_FOLDER'], 'pagenumbers.pdf')
    add_page_numbers(in_path, out_path)

    flash("Numéros de page ajoutés.")
    return send_file(out_path, as_attachment=True)


@app.route('/repair', methods=['POST'])
def route_repair():
    file = request.files.get('repair_file')
    if not file or file.filename.strip() == '':
        flash("Aucun fichier sélectionné pour la réparation.")
        return redirect(url_for('index'))

    if not is_pdf_file(file.filename):
        flash(f"Fichier invalide '{file.filename}' : veuillez sélectionner un PDF.")
        return redirect(url_for('index'))

    filename = secure_filename(file.filename)
    in_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(in_path)

    out_path = os.path.join(app.config['OUTPUT_FOLDER'], 'repaired.pdf')
    repair_pdf(in_path, out_path)

    flash("Réparation terminée.")
    return send_file(out_path, as_attachment=True)


@app.route('/unlock', methods=['POST'])
def route_unlock():
    file = request.files.get('unlock_file')
    password = request.form.get('unlock_password')

    if not file or file.filename.strip() == '':
        flash("Aucun fichier PDF pour le déverrouillage.")
        return redirect(url_for('index'))
    if not password:
        flash("Mot de passe manquant.")
        return redirect(url_for('index'))

    if not is_pdf_file(file.filename):
        flash(f"Fichier invalide '{file.filename}' : veuillez sélectionner un PDF.")

    filename = secure_filename(file.filename)
    in_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(in_path)

    out_path = os.path.join(app.config['OUTPUT_FOLDER'], 'unlocked.pdf')
    try:
        unlock_pdf(in_path, out_path, password)
        flash("PDF déverrouillé.")
        return send_file(out_path, as_attachment=True)
    except Exception as e:
        flash(f"Erreur lors du déverrouillage : {e}")
        return redirect(url_for('index'))


@app.route('/protect', methods=['POST'])
def route_protect():
    file = request.files.get('protect_file')
    password = request.form.get('protect_password')

    if not file or file.filename.strip() == '':
        flash("Aucun fichier sélectionné pour la protection.")
        return redirect(url_for('index'))
    if not password:
        flash("Mot de passe manquant.")
        return redirect(url_for('index'))

    if not is_pdf_file(file.filename):
        flash(f"Fichier invalide '{file.filename}' : veuillez sélectionner un PDF.")
        return redirect(url_for('index'))

    filename = secure_filename(file.filename)
    in_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(in_path)

    out_path = os.path.join(app.config['OUTPUT_FOLDER'], 'protected.pdf')
    protect_pdf(in_path, out_path, password)

    flash("PDF protégé.")
    return send_file(out_path, as_attachment=True)

@app.route('/images_to_pdf', methods=['POST'])
def route_images_to_pdf():
    files = request.files.getlist('images_files')

    if not files or len(files) == 0:
        flash("Aucune image sélectionnée pour la conversion en PDF.")
        return redirect(url_for('index'))

    saved_paths = []
    for f in files:
        if f.filename.strip() == '':
            continue

        filename = secure_filename(f.filename)
        f_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        f.save(f_path)
        saved_paths.append(f_path)

    if not saved_paths:
        flash("Veuillez sélectionner au moins une image.")
        return redirect(url_for('index'))

    output_file = os.path.join(app.config['OUTPUT_FOLDER'], 'converted.pdf')
    try:
        images_to_pdf(saved_paths, output_file)
        flash("Conversion Images → PDF terminée.")
        return send_file(output_file, as_attachment=True)
    except Exception as e:
        flash(f"Erreur conversion Images → PDF : {e}")
        return redirect(url_for('index'))

@app.route('/word_to_pdf', methods=['POST'])
def route_word_to_pdf():
    file = request.files.get('word_file')

    if not file or file.filename.strip() == '':
        flash("Aucun fichier Word sélectionné.")
        return redirect(url_for('index'))

    if not file.filename.lower().endswith('.docx'):
        flash("Fichier invalide : veuillez sélectionner un fichier Word (.docx).")
        return redirect(url_for('index'))

    filename = secure_filename(file.filename)
    in_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(in_path)

    out_path = os.path.join(app.config['OUTPUT_FOLDER'], filename.replace(".docx", ".pdf"))

    try:
        word_to_pdf(in_path, out_path)

        if not os.path.exists(out_path):
            flash("Erreur : le fichier PDF n'a pas été généré.")
            return redirect(url_for('index'))

        flash("Conversion Word → PDF terminée.")
        return send_file(out_path, as_attachment=True)

    except Exception as e:
        flash(f"Erreur conversion Word → PDF : {e}")
        return redirect(url_for('index'))

import shutil

@app.after_request
def cleanup_files(response):
    """ Supprime tous les fichiers temporaires après la requête """
    try:
        # Supprimer tous les fichiers du dossier UPLOAD_FOLDER et OUTPUT_FOLDER
        for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER']]:
            for filename in os.listdir(folder):
                file_path = os.path.join(folder, filename)
                if os.path.isfile(file_path):
                    os.remove(file_path)
    except Exception as e:
        print(f"⚠️ Erreur lors du nettoyage des fichiers : {e}")
    
    return response

# ================
# = Lancement
# ================
if __name__ == "__main__":
    app.run(debug=True)
